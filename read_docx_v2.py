
from docx import Document

doc = Document("/home/tsuyuri/Documents/web-compro-lembaga/Masukan 1(2).docx")

print("=== File Content ===\n")

for para_idx, paragraph in enumerate(doc.paragraphs):
    print(f"Paragraph {para_idx}:")
    for run_idx, run in enumerate(paragraph.runs):
        text = run.text.strip()
        if not text:
            continue
        
        # Check highlight color
        highlight = run.font.highlight_color
        highlight_str = str(highlight) if highlight is not None else "None"
        
        # Check text color
        color_str = "Auto"
        if run.font.color and run.font.color.rgb:
            color_str = str(run.font.color.rgb)
        
        print(f"  Run {run_idx}: '{text}'")
        print(f"    - Highlight: {highlight_str}")
        print(f"    - Text Color: {color_str}")
    print()

print("\n=== Tables ===\n")
for table_idx, table in enumerate(doc.tables):
    print(f"Table {table_idx}:")
    for row_idx, row in enumerate(table.rows):
        row_data = []
        for cell in row.cells:
            row_data.append(cell.text.strip())
        print(f"  Row {row_idx}: {row_data}")
